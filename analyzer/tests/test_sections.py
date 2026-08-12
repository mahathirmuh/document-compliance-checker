"""Per-section analysis (CLAUDE.md 7, 27 Phase 4)."""

from __future__ import annotations

from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.config import get_settings
from app.main import app
from app.parsers.base import SegmentKind, TextSegment
from app.services.analysis_service import get_analysis_service
from app.services.section_analyzer import UNTITLED_SECTION, SectionAnalyzer
from tests.conftest import CHINESE_TEXT, ENGLISH_TEXT, INDONESIAN_TEXT, build_docx


def heading(text: str) -> TextSegment:
    return TextSegment(text=text, kind=SegmentKind.HEADING, section=text)


def paragraph(text: str, section: str | None = None) -> TextSegment:
    return TextSegment(text=text, kind=SegmentKind.PARAGRAPH, section=section)


class TestSectionGrouping:
    @pytest.fixture
    def analyzer(self) -> SectionAnalyzer:
        return SectionAnalyzer()

    def test_a_heading_starts_a_new_section(self, analyzer: SectionAnalyzer) -> None:
        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            heading("2. Responsibility"),
            paragraph(ENGLISH_TEXT, "2. Responsibility"),
        ])

        assert [report.name for report in reports] == ["1. Scope", "2. Responsibility"]
        assert [report.sequence for report in reports] == [1, 2]

    def test_content_without_a_heading_falls_into_a_body_section(
        self, analyzer: SectionAnalyzer
    ) -> None:
        reports = analyzer.analyze([paragraph(ENGLISH_TEXT), paragraph(INDONESIAN_TEXT)])

        assert len(reports) == 1
        assert reports[0].name == UNTITLED_SECTION

    def test_it_reports_nothing_for_an_empty_document(self, analyzer: SectionAnalyzer) -> None:
        assert analyzer.analyze([]) == []

    def test_a_paginated_document_without_headings_falls_back_to_pages(
        self, analyzer: SectionAnalyzer
    ) -> None:
        # PDF is the dominant format in a real controlled-document library and
        # has no heading styles to read. Without this every PDF collapses into
        # one "(document body)" section and the per-section feature reports
        # nothing useful for exactly the documents that need it most.
        segments = [
            TextSegment(text=ENGLISH_TEXT, kind=SegmentKind.LINE, page=1),
            TextSegment(text=INDONESIAN_TEXT, kind=SegmentKind.LINE, page=1),
            TextSegment(text=ENGLISH_TEXT, kind=SegmentKind.LINE, page=2),
            TextSegment(text=CHINESE_TEXT, kind=SegmentKind.LINE, page=2),
        ]

        reports = analyzer.analyze(segments)

        assert [report.name for report in reports] == ["Page 1", "Page 2"]
        assert [report.page for report in reports] == [1, 2]

    def test_a_missing_language_is_located_to_a_page(self, analyzer: SectionAnalyzer) -> None:
        segments = [
            TextSegment(text=ENGLISH_TEXT, kind=SegmentKind.LINE, page=1),
            TextSegment(text=INDONESIAN_TEXT, kind=SegmentKind.LINE, page=1),
            TextSegment(text=CHINESE_TEXT, kind=SegmentKind.LINE, page=1),
            TextSegment(text=ENGLISH_TEXT, kind=SegmentKind.LINE, page=2),
            TextSegment(text=INDONESIAN_TEXT, kind=SegmentKind.LINE, page=2),
        ]

        reports = analyzer.analyze(segments)

        assert reports[0].missing == []
        assert reports[1].missing == ["zh"]
        assert reports[1].page == 2

    def test_headings_win_over_pages_when_a_document_has_both(
        self, analyzer: SectionAnalyzer
    ) -> None:
        # A heading is a more meaningful unit than a page boundary, which can
        # fall anywhere - including mid-sentence.
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope", page=1),
            TextSegment(text=ENGLISH_TEXT, kind=SegmentKind.PARAGRAPH, section="1. Scope", page=1),
            TextSegment(text=INDONESIAN_TEXT, kind=SegmentKind.PARAGRAPH, section="1. Scope", page=2),
        ]

        reports = analyzer.analyze(segments)

        assert [report.name for report in reports] == ["1. Scope"]


class TestSectionCoverage:
    @pytest.fixture
    def analyzer(self) -> SectionAnalyzer:
        return SectionAnalyzer()

    def test_a_fully_translated_section_reports_nothing_missing(
        self, analyzer: SectionAnalyzer
    ) -> None:
        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            paragraph(INDONESIAN_TEXT, "1. Scope"),
            paragraph(CHINESE_TEXT, "1. Scope"),
        ])

        assert reports[0].missing == []
        assert reports[0].short == []

    def test_it_locates_the_section_that_is_missing_a_language(
        self, analyzer: SectionAnalyzer
    ) -> None:
        # The whole point of Phase 4: not "this document is missing Mandarin"
        # but "section 2 is missing Mandarin".
        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            paragraph(INDONESIAN_TEXT, "1. Scope"),
            paragraph(CHINESE_TEXT, "1. Scope"),
            heading("2. Responsibility"),
            paragraph(ENGLISH_TEXT, "2. Responsibility"),
            paragraph(INDONESIAN_TEXT, "2. Responsibility"),
        ])

        assert reports[0].missing == []
        assert reports[1].missing == ["zh"]

    def test_each_language_is_counted_within_its_own_section(
        self, analyzer: SectionAnalyzer
    ) -> None:
        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            heading("2. Ruang Lingkup"),
            paragraph(INDONESIAN_TEXT, "2. Ruang Lingkup"),
        ])

        assert reports[0].profile.tallies["en"].characters > 0
        assert reports[0].profile.tallies["id"].characters == 0
        assert reports[1].profile.tallies["id"].characters > 0
        assert reports[1].profile.tallies["en"].characters == 0

    def test_a_short_section_is_measured_but_not_reported_against(self) -> None:
        # A heading or a two-line note cannot carry three languages, and
        # flagging it would bury the sections that matter.
        analyzer = SectionAnalyzer(min_section_chars=500)

        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph("Short English note only.", "1. Scope"),
        ])

        assert reports[0].total_characters > 0
        assert reports[0].missing == []

    def test_a_section_of_only_numbers_is_not_reported_as_missing_everything(
        self, analyzer: SectionAnalyzer
    ) -> None:
        reports = analyzer.analyze([
            heading("Appendix A"),
            paragraph("1234 5678 90-11 2026-01-01", "Appendix A"),
        ])

        # Digits carry no language, so the only meaningful text here is the
        # heading itself - far too little to expect three translations of.
        assert reports[0].total_characters < 20
        assert reports[0].missing == []
        assert reports[0].short == []


class TestShortTranslationDetection:
    def test_a_disproportionately_short_translation_is_flagged(self) -> None:
        analyzer = SectionAnalyzer(min_section_chars=50, short_translation_ratio=0.25)

        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            paragraph(INDONESIAN_TEXT, "1. Scope"),
            # Present, but a fraction of the others.
            paragraph("质量检验", "1. Scope"),
        ])

        assert "zh" in reports[0].short
        assert reports[0].missing == []

    def test_balanced_translations_are_not_flagged(self) -> None:
        analyzer = SectionAnalyzer(min_section_chars=50, short_translation_ratio=0.25)

        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            paragraph(INDONESIAN_TEXT, "1. Scope"),
            paragraph(CHINESE_TEXT, "1. Scope"),
        ])

        # Han is far denser per character than Latin; the ratio is generous
        # enough that density alone must not raise a finding.
        assert reports[0].short == []

    def test_a_missing_language_is_not_also_reported_as_short(self) -> None:
        analyzer = SectionAnalyzer(min_section_chars=50)

        reports = analyzer.analyze([
            heading("1. Scope"),
            paragraph(ENGLISH_TEXT, "1. Scope"),
            paragraph(INDONESIAN_TEXT, "1. Scope"),
        ])

        assert reports[0].missing == ["zh"]
        assert "zh" not in reports[0].short


class TestSectionsOverTheApi:
    @pytest.fixture
    def client(self, monkeypatch: pytest.MonkeyPatch) -> TestClient:
        monkeypatch.delenv("ANALYZER_API_KEY", raising=False)
        monkeypatch.delenv("ANALYZER_ALLOWED_ROOTS", raising=False)
        get_settings.cache_clear()
        get_analysis_service.cache_clear()

        with TestClient(app) as test_client:
            yield test_client

        get_analysis_service.cache_clear()

    @pytest.fixture
    def sectioned_docx(self, tmp_path: Path) -> Path:
        from docx import Document as DocxDocument

        path = tmp_path / "sectioned.docx"
        document = DocxDocument()

        document.add_heading("1. Scope", level=1)
        document.add_paragraph(ENGLISH_TEXT)
        document.add_paragraph(INDONESIAN_TEXT)
        document.add_paragraph(CHINESE_TEXT)

        document.add_heading("2. Responsibility", level=1)
        document.add_paragraph(ENGLISH_TEXT)
        document.add_paragraph(INDONESIAN_TEXT)

        document.save(str(path))
        return path

    def test_the_response_carries_a_section_breakdown(
        self, client: TestClient, sectioned_docx: Path
    ) -> None:
        body = client.post("/api/v1/analyze", json={"file_path": str(sectioned_docx)}).json()

        names = [section["name"] for section in body["sections"]]

        assert "1. Scope" in names
        assert "2. Responsibility" in names

    def test_a_missing_section_translation_is_raised_with_its_location(
        self, client: TestClient, sectioned_docx: Path
    ) -> None:
        body = client.post("/api/v1/analyze", json={"file_path": str(sectioned_docx)}).json()

        located = [
            issue for issue in body["issues"]
            if issue["type"] == "MISSING_SECTION_TRANSLATION"
        ]

        assert located, 'Expected the second section to be reported as missing Chinese.'
        assert located[0]["section"] == "2. Responsibility"
        assert located[0]["language"] == "zh"

    def test_a_fully_trilingual_document_raises_no_section_issues(
        self, client: TestClient, trilingual_complete_docx: Path
    ) -> None:
        body = client.post(
            "/api/v1/analyze", json={"file_path": str(trilingual_complete_docx)}
        ).json()

        assert [
            issue for issue in body["issues"]
            if issue["type"] == "MISSING_SECTION_TRANSLATION"
        ] == []

    def test_section_characters_are_reported_per_language(
        self, client: TestClient, sectioned_docx: Path
    ) -> None:
        body = client.post("/api/v1/analyze", json={"file_path": str(sectioned_docx)}).json()

        section = next(s for s in body["sections"] if s["name"] == "1. Scope")

        assert section["characters"]["en"] > 0
        assert section["characters"]["id"] > 0
        assert section["characters"]["zh"] > 0
        assert section["total_characters"] > 0

    def test_a_document_with_no_headings_still_reports_one_section(
        self, client: TestClient, tmp_path: Path
    ) -> None:
        path = build_docx(tmp_path / "flat.docx", [ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT])

        body = client.post("/api/v1/analyze", json={"file_path": str(path)}).json()

        assert len(body["sections"]) == 1
        assert body["sections"][0]["name"] == UNTITLED_SECTION
