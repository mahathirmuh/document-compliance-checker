"""Shared fixtures.

Documents are built at test time rather than committed as binaries. That keeps
what each fixture actually contains visible in the diff - "this one has 40
Chinese characters" is a fact you can read, not one you have to open Word to
discover - and it keeps the repository free of opaque files (CLAUDE.md 29).
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from app.config import get_settings
from app.detectors.language_detector import _build_detector

# --------------------------------------------------------------------------- #
# Sample text. Long enough to be classified reliably and to clear the default
# thresholds (EN/ID 100 characters, ZH 50).
# --------------------------------------------------------------------------- #

ENGLISH_TEXT = (
    "This standard operating procedure describes how finished goods are inspected "
    "before they are released from the warehouse. The quality inspector must record "
    "every deviation in the inspection log and notify the shift supervisor immediately. "
    "No shipment may leave the site without a signed release form."
)

INDONESIAN_TEXT = (
    "Prosedur operasional standar ini menjelaskan bagaimana barang jadi diperiksa "
    "sebelum dikeluarkan dari gudang. Petugas mutu wajib mencatat setiap penyimpangan "
    "pada buku catatan pemeriksaan dan segera memberitahu penyelia. Tidak ada pengiriman "
    "yang boleh meninggalkan lokasi tanpa formulir pelepasan yang telah ditandatangani."
)

CHINESE_TEXT = (
    "本标准操作程序说明成品在出库之前应如何检验。质量检验员必须在检验记录中登记每一项偏差，"
    "并立即通知当班主管。未经签署放行单，任何货物都不得离开厂区。"
)

#: Below the default 50-character Chinese threshold - used to prove that a
#: token translation does not count as coverage.
CHINESE_TOKEN = "质量"

JAPANESE_TEXT = "これは日本語の文書です。品質管理の手順を説明します。"


@pytest.fixture(autouse=True)
def _reset_settings_cache():
    """Stop configuration leaking between tests."""
    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture(scope="session", autouse=True)
def _warm_language_models():
    """Load lingua once for the whole session rather than per test."""
    _build_detector()


# --------------------------------------------------------------------------- #
# Builders
# --------------------------------------------------------------------------- #


def build_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> Path:
    document = DocxDocument()

    for text in paragraphs:
        document.add_paragraph(text)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value

    document.save(str(path))
    return path


def build_xlsx(path: Path, rows: list[list[str]], sheet_title: str = "Procedure") -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_title

    for row in rows:
        sheet.append(row)

    workbook.save(str(path))
    return path


def build_pdf(path: Path, lines: list[str]) -> Path:
    """A real PDF with a real text layer.

    Only Latin text: the default reportlab fonts have no CJK glyphs, so
    Chinese coverage is exercised through the DOCX, XLSX and TXT fixtures
    instead.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=A4)
    y = 800

    for line in lines:
        pdf.drawString(60, y, line)
        y -= 16
        if y < 60:
            pdf.showPage()
            y = 800

    pdf.save()
    return path


def build_image_only_pdf(path: Path) -> Path:
    """A PDF with a drawn rectangle and no text layer - i.e. a stand-in scan."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.rect(72, 500, 400, 200, fill=1)
    pdf.showPage()
    pdf.save()
    return path


# --------------------------------------------------------------------------- #
# The CLAUDE.md 29 fixture set
# --------------------------------------------------------------------------- #


@pytest.fixture
def english_only_docx(tmp_path: Path) -> Path:
    return build_docx(tmp_path / "english_only.docx", [ENGLISH_TEXT])


@pytest.fixture
def indonesian_only_docx(tmp_path: Path) -> Path:
    return build_docx(tmp_path / "indonesian_only.docx", [INDONESIAN_TEXT])


@pytest.fixture
def mandarin_only_docx(tmp_path: Path) -> Path:
    return build_docx(tmp_path / "mandarin_only.docx", [CHINESE_TEXT])


@pytest.fixture
def trilingual_complete_docx(tmp_path: Path) -> Path:
    return build_docx(
        tmp_path / "trilingual_complete.docx",
        [ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT],
    )


@pytest.fixture
def trilingual_partial_docx(tmp_path: Path) -> Path:
    """All three languages present, but Chinese is only a token."""
    return build_docx(
        tmp_path / "trilingual_partial.docx",
        [ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TOKEN],
    )


@pytest.fixture
def trilingual_table_docx(tmp_path: Path) -> Path:
    """The layout these documents actually use: one language per column."""
    return build_docx(
        tmp_path / "trilingual_table.docx",
        ["Inspection Procedure"],
        table_rows=[
            ["English", "Indonesia", "中文"],
            [ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT],
        ],
    )


@pytest.fixture
def empty_document_docx(tmp_path: Path) -> Path:
    return build_docx(tmp_path / "empty_document.docx", [])


@pytest.fixture
def large_document_docx(tmp_path: Path) -> Path:
    paragraphs: list[str] = []
    for _ in range(200):
        paragraphs.extend([ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT])

    return build_docx(tmp_path / "large_document.docx", paragraphs)


@pytest.fixture
def scanned_document_pdf(tmp_path: Path) -> Path:
    return build_image_only_pdf(tmp_path / "scanned_document.pdf")


@pytest.fixture
def bilingual_pdf(tmp_path: Path) -> Path:
    return build_pdf(
        tmp_path / "bilingual.pdf",
        [ENGLISH_TEXT[:90], ENGLISH_TEXT[90:180], INDONESIAN_TEXT[:90], INDONESIAN_TEXT[90:180]],
    )


@pytest.fixture
def trilingual_xlsx(tmp_path: Path) -> Path:
    return build_xlsx(
        tmp_path / "trilingual.xlsx",
        [
            ["Section", "English", "Indonesia", "中文"],
            ["Scope", ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT],
        ],
    )


@pytest.fixture
def trilingual_txt(tmp_path: Path) -> Path:
    path = tmp_path / "trilingual.txt"
    path.write_text(
        "\n".join([ENGLISH_TEXT, INDONESIAN_TEXT, CHINESE_TEXT]),
        encoding="utf-8",
    )
    return path
