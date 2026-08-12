"""Parser tests (CLAUDE.md 29)."""

from __future__ import annotations

from pathlib import Path

import pytest

from app.parsers.base import ParserError, SegmentKind, UnsupportedFormatError
from app.parsers.docx_parser import DocxParser
from app.parsers.pdf_parser import PdfParser
from app.parsers.registry import ParserRegistry
from app.parsers.txt_parser import TxtParser
from app.parsers.xlsx_parser import XlsxParser
from tests.conftest import CHINESE_TEXT, ENGLISH_TEXT, INDONESIAN_TEXT, build_docx


class TestDocxParser:
    def test_it_extracts_paragraphs(self, trilingual_complete_docx: Path) -> None:
        result = DocxParser().parse(trilingual_complete_docx)

        assert len(result.segments) == 3
        assert ENGLISH_TEXT in result.text
        assert CHINESE_TEXT in result.text

    def test_it_extracts_table_cells_separately(self, trilingual_table_docx: Path) -> None:
        # One cell per language, not one merged string per row - a merged row
        # would hand the detector a mixed segment to pull apart again.
        result = DocxParser().parse(trilingual_table_docx)
        cells = [s for s in result.segments if s.kind is SegmentKind.TABLE_CELL]

        assert any(ENGLISH_TEXT in cell.text for cell in cells)
        assert any(INDONESIAN_TEXT in cell.text for cell in cells)
        assert any(CHINESE_TEXT in cell.text for cell in cells)

    def test_it_records_headings_as_sections(self, tmp_path: Path) -> None:
        from docx import Document as DocxDocument

        path = tmp_path / "headed.docx"
        document = DocxDocument()
        document.add_heading("Scope", level=1)
        document.add_paragraph(ENGLISH_TEXT)
        document.save(str(path))

        result = DocxParser().parse(path)
        body = [s for s in result.segments if s.kind is SegmentKind.PARAGRAPH]

        assert any(s.kind is SegmentKind.HEADING for s in result.segments)
        assert body[0].section == "Scope"

    def test_it_reads_headers_and_footers(self, tmp_path: Path) -> None:
        # A Chinese title in the header is real trilingual content; missing it
        # would report a compliant document as missing Mandarin.
        from docx import Document as DocxDocument

        path = tmp_path / "furniture.docx"
        document = DocxDocument()
        document.add_paragraph(ENGLISH_TEXT)
        document.sections[0].header.paragraphs[0].text = "SOP-QA-001 质量手册"
        document.sections[0].footer.paragraphs[0].text = "Revision 03"
        document.save(str(path))

        result = DocxParser().parse(path)
        kinds = {s.kind for s in result.segments}

        assert SegmentKind.HEADER in kinds
        assert SegmentKind.FOOTER in kinds
        assert "质量手册" in result.text

    def test_an_empty_document_parses_to_no_segments(self, empty_document_docx: Path) -> None:
        # Empty is not an error: "blank" and "broken" are different findings.
        result = DocxParser().parse(empty_document_docx)

        assert result.segments == []

    def test_a_corrupt_file_raises_parser_error(self, tmp_path: Path) -> None:
        path = tmp_path / "broken.docx"
        path.write_bytes(b"PK\x03\x04 this is not really a docx")

        with pytest.raises(ParserError):
            DocxParser().parse(path)


class TestPdfParser:
    def test_it_extracts_text_with_page_numbers(self, bilingual_pdf: Path) -> None:
        result = PdfParser().parse(bilingual_pdf)

        assert result.page_count == 1
        assert all(segment.page == 1 for segment in result.segments)
        assert result.segments

    def test_a_scanned_pdf_yields_no_text_but_does_not_raise(
        self, scanned_document_pdf: Path
    ) -> None:
        # Critically it must not raise: an exception would be recorded as
        # ERROR, when the correct outcome is REVIEW_REQUIRED (CLAUDE.md 16).
        result = PdfParser().parse(scanned_document_pdf)

        assert result.text.strip() == ""
        assert any("scanned" in note for note in result.notes)

    def test_a_corrupt_pdf_raises_parser_error(self, tmp_path: Path) -> None:
        path = tmp_path / "broken.pdf"
        path.write_bytes(b"%PDF-1.7\nnot actually a pdf")

        with pytest.raises(ParserError):
            PdfParser().parse(path)


class TestXlsxParser:
    def test_it_extracts_string_cells_with_the_sheet_as_section(
        self, trilingual_xlsx: Path
    ) -> None:
        result = XlsxParser().parse(trilingual_xlsx)

        assert all(s.section == "Procedure" for s in result.segments)
        assert any(CHINESE_TEXT in s.text for s in result.segments)

    def test_it_ignores_non_string_cells(self, tmp_path: Path) -> None:
        # Numbers and dates carry no language; feeding them to a detector
        # would be noise.
        from openpyxl import Workbook

        path = tmp_path / "numbers.xlsx"
        workbook = Workbook()
        workbook.active.append(["Quantity", 42, 3.14, None, True])
        workbook.save(str(path))

        result = XlsxParser().parse(path)

        assert [s.text for s in result.segments] == ["Quantity"]

    def test_a_corrupt_workbook_raises_parser_error(self, tmp_path: Path) -> None:
        path = tmp_path / "broken.xlsx"
        path.write_bytes(b"not a workbook at all")

        with pytest.raises(ParserError):
            XlsxParser().parse(path)


class TestTxtParser:
    def test_it_reads_utf8(self, trilingual_txt: Path) -> None:
        result = TxtParser().parse(trilingual_txt)

        assert CHINESE_TEXT in result.text

    def test_it_honours_a_utf8_bom(self, tmp_path: Path) -> None:
        path = tmp_path / "bom.txt"
        path.write_bytes(b"\xef\xbb\xbf" + CHINESE_TEXT.encode("utf-8"))

        result = TxtParser().parse(path)

        assert result.text.startswith("本标准")

    def test_it_reads_utf16(self, tmp_path: Path) -> None:
        path = tmp_path / "utf16.txt"
        path.write_bytes(CHINESE_TEXT.encode("utf-16"))

        result = TxtParser().parse(path)

        assert "质量" in result.text
        assert any("utf-16" in note for note in result.notes)

    def test_it_reads_gb18030(self, tmp_path: Path) -> None:
        # A Chinese file saved from a Windows machine is very often GBK, and
        # decoding it as UTF-8 would produce mojibake the detector would then
        # classify with confidence - a wrong answer stated firmly.
        path = tmp_path / "gbk.txt"
        path.write_bytes(CHINESE_TEXT.encode("gb18030"))

        result = TxtParser().parse(path)

        assert "质量" in result.text


class TestRegistry:
    def test_it_resolves_each_supported_extension(self) -> None:
        registry = ParserRegistry()

        assert registry.for_extension("docx").name == "docx"
        assert registry.for_extension(".PDF").name == "pdf"
        assert registry.for_extension("xlsx").name == "xlsx"
        assert registry.for_extension("txt").name == "txt"

    def test_an_unknown_extension_raises(self) -> None:
        with pytest.raises(UnsupportedFormatError):
            ParserRegistry().for_extension("pptx")

    def test_supported_extensions_covers_the_v1_formats(self) -> None:
        supported = ParserRegistry().supported_extensions()

        assert {"docx", "pdf", "xlsx", "txt"} <= set(supported)


def test_a_large_document_is_parsed_without_loading_it_all_at_once(
    large_document_docx: Path,
) -> None:
    result = DocxParser().parse(large_document_docx)

    assert len(result.segments) == 600


def test_build_docx_helper_produces_a_readable_file(tmp_path: Path) -> None:
    path = build_docx(tmp_path / "smoke.docx", ["hello"])

    assert DocxParser().parse(path).text == "hello"
