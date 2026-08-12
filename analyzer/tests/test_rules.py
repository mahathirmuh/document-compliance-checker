"""Document Control rules (CLAUDE.md 7, 27 Phase 5)."""

from __future__ import annotations

from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.config import get_settings
from app.detectors.language_detector import LanguageDetector
from app.main import app
from app.parsers.base import ExtractedDocument, SegmentKind, TextSegment
from app.rules.base import RuleContext
from app.rules.document_code import DocumentCodePlacementRule
from app.rules.formatting import FontColorRule
from app.rules.language_order import LanguageOrderRule
from app.rules.registry import RuleRegistry
from app.rules.structure import CoverPageRule, HeaderFooterRule
from app.services.analysis_service import get_analysis_service
from app.services.section_analyzer import SectionAnalyzer
from tests.conftest import CHINESE_TEXT, ENGLISH_TEXT, INDONESIAN_TEXT


def context_for(
    segments: list[TextSegment],
    config: dict | None = None,
    parser: str = "docx",
    page_count: int | None = None,
) -> RuleContext:
    document = ExtractedDocument(segments=segments, parser=parser, page_count=page_count)
    detector = LanguageDetector()

    return RuleContext(
        document=document,
        sections=SectionAnalyzer(detector).analyze(segments),
        detector=detector,
        config=config or {},
    )


class TestLanguageOrder:
    def test_the_expected_order_passes(self) -> None:
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope"),
            TextSegment(text=ENGLISH_TEXT, section="1. Scope"),
            TextSegment(text=INDONESIAN_TEXT, section="1. Scope"),
            TextSegment(text=CHINESE_TEXT, section="1. Scope"),
        ]

        outcome = LanguageOrderRule().evaluate(context_for(segments, {"order": ["en", "id", "zh"]}))

        assert outcome.passed

    def test_a_reversed_order_is_reported_with_its_section(self) -> None:
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope"),
            TextSegment(text=CHINESE_TEXT, section="1. Scope"),
            TextSegment(text=INDONESIAN_TEXT, section="1. Scope"),
            TextSegment(text=ENGLISH_TEXT, section="1. Scope"),
        ]

        outcome = LanguageOrderRule().evaluate(context_for(segments, {"order": ["en", "id", "zh"]}))

        assert not outcome.passed
        finding = outcome.findings[0]
        assert finding.issue_type == "WRONG_LANGUAGE_ORDER"
        assert finding.section == "1. Scope"
        assert finding.metadata["observed"] == ["zh", "id", "en"]

    def test_a_section_with_one_language_has_no_order_to_get_wrong(self) -> None:
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope"),
            TextSegment(text=ENGLISH_TEXT, section="1. Scope"),
        ]

        assert LanguageOrderRule().evaluate(context_for(segments)).passed

    def test_a_missing_language_is_not_treated_as_wrong_order(self) -> None:
        # A section that legitimately covers only English and Chinese must not
        # fail for "missing" Indonesian - that is the translation rule's job.
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope"),
            TextSegment(text=ENGLISH_TEXT, section="1. Scope"),
            TextSegment(text=CHINESE_TEXT, section="1. Scope"),
        ]

        assert LanguageOrderRule().evaluate(context_for(segments)).passed

    def test_a_custom_order_is_honoured(self) -> None:
        segments = [
            TextSegment(text="1. Scope", kind=SegmentKind.HEADING, section="1. Scope"),
            TextSegment(text=CHINESE_TEXT, section="1. Scope"),
            TextSegment(text=ENGLISH_TEXT, section="1. Scope"),
        ]

        outcome = LanguageOrderRule().evaluate(context_for(segments, {"order": ["zh", "en", "id"]}))

        assert outcome.passed


class TestDocumentCode:
    def test_it_extracts_a_code_and_revision_from_the_header(self) -> None:
        segments = [
            TextSegment(text="MTI-ENV-EVM-SOP-002 Rev. 006", kind=SegmentKind.HEADER),
            TextSegment(text=ENGLISH_TEXT),
        ]

        rule = DocumentCodePlacementRule()
        outcome = rule.evaluate(context_for(segments))

        assert outcome.passed
        assert rule.extracted["document_code"] == "MTI-ENV-EVM-SOP-002"
        assert rule.extracted["revision"] == "006"

    def test_it_reads_the_underscore_form_seen_on_real_documents(self) -> None:
        # Regression: real headers read "MTI-ENV-EVM-SOP-002_Rev. 006", and
        # "_" is a word character - so a \b boundary silently fails on both
        # the code and the revision. Caught only by running the real files.
        segments = [TextSegment(text="MTI-ENV-EVM-SOP-002_Rev. 006", page=1)]

        rule = DocumentCodePlacementRule()
        outcome = rule.evaluate(context_for(segments, parser="pdf", page_count=5))

        assert rule.extracted["document_code"] == "MTI-ENV-EVM-SOP-002"
        assert rule.extracted["revision"] == "006"
        assert outcome.passed

    def test_a_missing_code_is_reported(self) -> None:
        segments = [TextSegment(text=ENGLISH_TEXT)]

        outcome = DocumentCodePlacementRule().evaluate(context_for(segments))

        types = {finding.issue_type for finding in outcome.findings}
        assert "MISSING_DOCUMENT_CODE" in types
        assert "MISSING_REVISION" in types

    def test_the_body_is_not_searched_for_the_code(self) -> None:
        # Scanning the whole document would match a code quoted in a
        # cross-reference on page 40 and report a blank header as compliant.
        segments = [
            TextSegment(text="Introduction", page=1),
            *[TextSegment(text=ENGLISH_TEXT, page=page) for page in range(2, 8)],
            TextSegment(text="See also MTI-ENV-EVM-SOP-002 for details.", page=8),
        ]

        rule = DocumentCodePlacementRule()
        rule.evaluate(context_for(segments, parser="pdf", page_count=8))

        assert rule.extracted["document_code"] is None

    def test_a_custom_pattern_is_honoured(self) -> None:
        segments = [TextSegment(text="Doc ref: QA/2026/0042", kind=SegmentKind.HEADER)]

        rule = DocumentCodePlacementRule()
        rule.evaluate(
            context_for(
                segments,
                {"code_pattern": r"[A-Z]{2}/\d{4}/\d{4}", "require_revision": False},
            )
        )

        assert rule.extracted["document_code"] == "QA/2026/0042"

    def test_an_invalid_pattern_does_not_fail_the_document(self) -> None:
        # A bad regex is a configuration mistake, not a document problem.
        segments = [TextSegment(text="MTI-ENV-EVM-SOP-002", kind=SegmentKind.HEADER)]

        rule = DocumentCodePlacementRule()
        rule.evaluate(context_for(segments, {"code_pattern": "([unclosed"}))

        assert rule.extracted["document_code"] is None

    def test_extraction_does_not_leak_between_documents(self) -> None:
        registry = RuleRegistry()
        config = {"document_code": {"enabled": True, "require_revision": False}}
        detector = LanguageDetector()

        first = ExtractedDocument(
            segments=[TextSegment(text="MTI-ENV-EVM-SOP-002", kind=SegmentKind.HEADER)],
            parser="docx",
        )
        _, found = registry.run(config, first, [], detector)
        assert found["document_code"] == "MTI-ENV-EVM-SOP-002"

        second = ExtractedDocument(
            segments=[TextSegment(text="no code here at all")], parser="docx"
        )
        _, found_again = registry.run(config, second, [], detector)

        assert found_again.get("document_code") is None


class TestHeaderFooter:
    def test_a_docx_with_both_passes(self) -> None:
        segments = [
            TextSegment(text="MTI-ENV-EVM-SOP-002", kind=SegmentKind.HEADER),
            TextSegment(text=ENGLISH_TEXT),
            TextSegment(text="Page 1 of 12", kind=SegmentKind.FOOTER),
        ]

        assert HeaderFooterRule().evaluate(context_for(segments)).passed

    def test_a_docx_missing_a_footer_is_reported(self) -> None:
        segments = [
            TextSegment(text="MTI-ENV-EVM-SOP-002", kind=SegmentKind.HEADER),
            TextSegment(text=ENGLISH_TEXT),
        ]

        outcome = HeaderFooterRule().evaluate(context_for(segments))

        assert {f.issue_type for f in outcome.findings} == {"MISSING_FOOTER"}

    def test_a_short_pdf_reports_that_it_could_not_tell(self) -> None:
        # With two pages, "appears on most pages" is not a meaningful
        # statement - and a guess here would be worse than an honest skip.
        segments = [
            TextSegment(text=ENGLISH_TEXT, page=1),
            TextSegment(text=ENGLISH_TEXT, page=2),
        ]

        outcome = HeaderFooterRule().evaluate(context_for(segments, parser="pdf", page_count=2))

        assert outcome.applicable is False
        assert "cannot be distinguished" in (outcome.skipped_reason or "")

    def test_a_pdf_header_is_inferred_from_repetition(self) -> None:
        segments = []
        for page in range(1, 7):
            segments.append(TextSegment(text="MTI-ENV-EVM-SOP-002 Rev. 006", page=page))
            segments.append(TextSegment(text=ENGLISH_TEXT, page=page))
            segments.append(TextSegment(text=f"Page {page} of 6", page=page))

        outcome = HeaderFooterRule().evaluate(context_for(segments, parser="pdf", page_count=6))

        assert outcome.applicable is True
        assert outcome.passed


class TestCoverPage:
    def test_an_opening_with_a_code_and_title_passes(self) -> None:
        segments = [
            TextSegment(text="MTI-ENV-EVM-SOP-002", page=1),
            TextSegment(text="Hazardous Waste Management and Handling", page=1),
        ]

        assert CoverPageRule().evaluate(context_for(segments, parser="pdf", page_count=3)).passed

    def test_an_opening_without_a_code_is_reported(self) -> None:
        segments = [TextSegment(text="Hazardous Waste Management and Handling", page=1)]

        outcome = CoverPageRule().evaluate(context_for(segments, parser="pdf", page_count=3))

        assert not outcome.passed
        assert outcome.findings[0].issue_type == "INVALID_COVER_PAGE"


class TestFontColor:
    def test_permitted_colours_pass(self) -> None:
        segments = [
            TextSegment(text=ENGLISH_TEXT, font_colors=frozenset({"000000"})),
            TextSegment(text=INDONESIAN_TEXT, font_colors=frozenset({"1F497D"})),
        ]

        outcome = FontColorRule().evaluate(
            context_for(segments, {"allowed": ["000000", "1F497D"]})
        )

        assert outcome.passed

    def test_an_unexpected_colour_is_reported_once_per_colour(self) -> None:
        segments = [
            TextSegment(text="first offender", font_colors=frozenset({"FF0000"})),
            TextSegment(text="second offender", font_colors=frozenset({"FF0000"})),
        ]

        outcome = FontColorRule().evaluate(context_for(segments, {"allowed": ["000000"]}))

        assert len(outcome.findings) == 1
        assert outcome.findings[0].metadata["occurrences"] == 2

    def test_a_pdf_is_reported_as_not_checked_rather_than_passing(self) -> None:
        # The most important behaviour in this rule. Calling an unreadable
        # format compliant would hand a clean bill of health to exactly the
        # documents least likely to deserve one.
        segments = [TextSegment(text=ENGLISH_TEXT, page=1)]

        outcome = FontColorRule().evaluate(context_for(segments, {"allowed": ["000000"]}, parser="pdf"))

        assert outcome.applicable is False
        assert outcome.passed is False
        assert "cannot be read from a PDF" in (outcome.skipped_reason or "")

    def test_a_themed_docx_is_also_reported_as_not_checked(self) -> None:
        # Text inheriting its colour from the template reports nothing, which
        # must not be mistaken for "uses only black".
        segments = [TextSegment(text=ENGLISH_TEXT)]

        outcome = FontColorRule().evaluate(context_for(segments, {"allowed": ["000000"]}))

        assert outcome.applicable is False
        assert "inherits" in (outcome.skipped_reason or "")


class TestRegistry:
    def test_rules_are_off_unless_enabled(self) -> None:
        registry = RuleRegistry()
        document = ExtractedDocument(segments=[TextSegment(text=ENGLISH_TEXT)], parser="docx")

        outcomes, _ = registry.run(None, document, [], LanguageDetector())

        assert outcomes == []

    def test_only_enabled_rules_run(self) -> None:
        registry = RuleRegistry()
        document = ExtractedDocument(segments=[TextSegment(text=ENGLISH_TEXT)], parser="docx")

        outcomes, _ = registry.run(
            {"document_code": {"enabled": True}, "font_color": {"enabled": False}},
            document,
            [],
            LanguageDetector(),
        )

        assert [outcome.rule for outcome in outcomes] == ["document_code"]

    def test_every_rule_is_listed(self) -> None:
        assert set(RuleRegistry().available()) == {
            "language_order",
            "document_code",
            "header_footer",
            "cover_page",
            "font_color",
        }


class TestRulesOverTheApi:
    @pytest.fixture
    def client(self, monkeypatch: pytest.MonkeyPatch) -> TestClient:
        monkeypatch.delenv("ANALYZER_API_KEY", raising=False)
        monkeypatch.delenv("ANALYZER_ALLOWED_ROOTS", raising=False)
        get_settings.cache_clear()
        get_analysis_service.cache_clear()

        with TestClient(app) as test_client:
            yield test_client

        get_analysis_service.cache_clear()

    @pytest.fixture
    def coded_docx(self, tmp_path: Path) -> Path:
        from docx import Document as DocxDocument

        path = tmp_path / "coded.docx"
        document = DocxDocument()
        document.sections[0].header.paragraphs[0].text = "MTI-ACP-APM-SOP-001 Rev. 000"
        document.add_heading("1. Scope", level=1)
        document.add_paragraph(ENGLISH_TEXT)
        document.add_paragraph(INDONESIAN_TEXT)
        document.add_paragraph(CHINESE_TEXT)
        document.save(str(path))

        return path

    def test_no_rules_run_when_none_are_requested(
        self, client: TestClient, coded_docx: Path
    ) -> None:
        body = client.post("/api/v1/analyze", json={"file_path": str(coded_docx)}).json()

        assert body["rules"] == []

    def test_enabled_rules_are_reported(self, client: TestClient, coded_docx: Path) -> None:
        body = client.post(
            "/api/v1/analyze",
            json={
                "file_path": str(coded_docx),
                "rules": {
                    "document_code": {"enabled": True},
                    "language_order": {"enabled": True, "order": ["en", "id", "zh"]},
                },
            },
        ).json()

        outcomes = {rule["rule"]: rule for rule in body["rules"]}

        assert outcomes["document_code"]["passed"] is True
        assert outcomes["language_order"]["passed"] is True

    def test_the_extracted_code_and_revision_come_back(
        self, client: TestClient, coded_docx: Path
    ) -> None:
        body = client.post(
            "/api/v1/analyze",
            json={"file_path": str(coded_docx), "rules": {"document_code": {"enabled": True}}},
        ).json()

        assert body["metadata"]["document_code"] == "MTI-ACP-APM-SOP-001"
        assert body["metadata"]["revision"] == "000"

    def test_a_skipped_rule_says_why(self, client: TestClient, bilingual_pdf: Path) -> None:
        body = client.post(
            "/api/v1/analyze",
            json={"file_path": str(bilingual_pdf), "rules": {"font_color": {"enabled": True}}},
        ).json()

        font_rule = next(rule for rule in body["rules"] if rule["rule"] == "font_color")

        assert font_rule["applicable"] is False
        assert font_rule["passed"] is False
        assert "PDF" in font_rule["skipped_reason"]

    def test_rule_findings_appear_as_located_issues(
        self, client: TestClient, tmp_path: Path
    ) -> None:
        from docx import Document as DocxDocument

        path = tmp_path / "reversed.docx"
        document = DocxDocument()
        document.add_heading("1. Scope", level=1)
        document.add_paragraph(CHINESE_TEXT)
        document.add_paragraph(ENGLISH_TEXT)
        document.save(str(path))

        body = client.post(
            "/api/v1/analyze",
            json={
                "file_path": str(path),
                "rules": {"language_order": {"enabled": True, "order": ["en", "id", "zh"]}},
            },
        ).json()

        ordering = [i for i in body["issues"] if i["type"] == "WRONG_LANGUAGE_ORDER"]

        assert ordering
        assert ordering[0]["section"] == "1. Scope"
        assert ordering[0]["metadata"]["rule"] == "language_order"
