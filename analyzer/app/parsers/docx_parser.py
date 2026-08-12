"""DOCX extraction.

A controlled document keeps a lot of its trilingual content outside the main
body flow: bilingual tables, a Chinese title in the header, a document code in
the footer. Reading only `document.paragraphs` would miss all of it and report
a perfectly compliant SOP as missing Mandarin, so tables, headers and footers
are all walked.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.parsers.base import (
    DocumentParser,
    ExtractedDocument,
    ParserError,
    SegmentKind,
    TextSegment,
)


class DocxParser(DocumentParser):
    extensions = ("docx",)
    name = "docx"

    def parse(self, path: Path) -> ExtractedDocument:
        try:
            document = DocxDocument(str(path))
        except Exception as exc:
            raise ParserError("The Word document could not be opened. It may be corrupt.") from exc

        result = ExtractedDocument(parser=self.name)
        current_section: str | None = None

        try:
            for block in self._iter_body(document):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if not text:
                        continue

                    if self._is_heading(block):
                        current_section = text
                        result.segments.append(
                            TextSegment(text=text, kind=SegmentKind.HEADING, section=text)
                        )
                    else:
                        result.segments.append(
                            TextSegment(
                                text=text,
                                kind=SegmentKind.PARAGRAPH,
                                section=current_section,
                            )
                        )

                elif isinstance(block, Table):
                    result.segments.extend(self._table_segments(block, current_section))

            result.segments.extend(self._section_furniture(document))
        except Exception as exc:
            raise ParserError("The Word document could not be read to the end.") from exc

        return result

    # ------------------------------------------------------------------ #

    def _iter_body(self, document: DocxDocumentType):
        """Yield paragraphs and tables in the order they appear.

        python-docx exposes `paragraphs` and `tables` as separate flat lists,
        which loses their interleaving - and with it the ability to attribute
        a table to the heading above it. Walking the XML body keeps that.
        """
        body = document.element.body

        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, document)
            elif child.tag == qn("w:tbl"):
                yield Table(child, document)

    def _is_heading(self, paragraph: Paragraph) -> bool:
        style = (paragraph.style.name or "") if paragraph.style else ""
        return style.startswith("Heading") or style in {"Title", "Subtitle"}

    def _table_segments(self, table: Table, section: str | None) -> list[TextSegment]:
        """One segment per cell.

        Per cell rather than per row: a bilingual table usually puts each
        language in its own column, and merging a row would produce a mixed
        string that the detector then has to pull apart again.
        """
        segments: list[TextSegment] = []

        # Merged cells appear once per grid position they span, so they have
        # to be deduplicated by their underlying element.
        #
        # The references are held in a list rather than only their ids in a
        # set: python-docx builds the lxml proxy on each access, and a proxy
        # that gets collected frees its id for reuse - which would silently
        # skip a *different* cell later in the table.
        seen_elements: list[object] = []
        seen_ids: set[int] = set()

        for row in table.rows:
            for cell in row.cells:
                element = cell._tc

                if id(element) in seen_ids:
                    continue

                seen_elements.append(element)
                seen_ids.add(id(element))

                text = cell.text.strip()
                if text:
                    segments.append(
                        TextSegment(text=text, kind=SegmentKind.TABLE_CELL, section=section)
                    )

        return segments

    def _section_furniture(self, document: DocxDocumentType) -> list[TextSegment]:
        """Headers and footers from every section."""
        segments: list[TextSegment] = []

        for section in document.sections:
            for part, kind in (
                (section.header, SegmentKind.HEADER),
                (section.footer, SegmentKind.FOOTER),
            ):
                if part is None:
                    continue

                for paragraph in part.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        segments.append(TextSegment(text=text, kind=kind))

                for table in part.tables:
                    segments.extend(self._table_segments(table, section=None))

        return segments
